"""Generate ATS-friendly DOCX resumes by cloning Resume_AI.docx formatting exactly."""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from copy import deepcopy


def clone_paragraph_format(source_para, target_para):
    """Clone all formatting from source paragraph to target paragraph."""
    # Copy paragraph format
    pf_source = source_para.paragraph_format
    pf_target = target_para.paragraph_format

    # Copy alignment
    pf_target.alignment = pf_source.alignment

    # Copy indentation
    pf_target.left_indent = pf_source.left_indent
    pf_target.right_indent = pf_source.right_indent
    pf_target.first_line_indent = pf_source.first_line_indent

    # Copy spacing
    pf_target.space_before = pf_source.space_before
    pf_target.space_after = pf_source.space_after

    # Copy line spacing
    pf_target.line_spacing = pf_source.line_spacing
    pf_target.line_spacing_rule = pf_source.line_spacing_rule

    # Copy other properties
    if pf_source.keep_together is not None:
        pf_target.keep_together = pf_source.keep_together
    if pf_source.keep_with_next is not None:
        pf_target.keep_with_next = pf_source.keep_with_next
    if pf_source.page_break_before is not None:
        pf_target.page_break_before = pf_source.page_break_before
    if pf_source.widow_control is not None:
        pf_target.widow_control = pf_source.widow_control


def clone_run_format(source_run, target_run):
    """Clone all formatting from source run to target run."""
    # Copy font properties
    source_font = source_run.font
    target_font = target_run.font

    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline
    target_font.color.rgb = source_font.color.rgb if source_font.color.rgb else None


def create_resume_docx(content: str, output_path: str, template_path: str = None) -> str:
    """
    Create resume DOCX by cloning Resume_AI.docx paragraph formats exactly.

    This reads Resume_AI.docx to extract exact formatting for:
    - Name (paragraph 0)
    - Contact (paragraph 1)
    - Section headers (paragraph 3, 7)
    - Job titles (paragraph 9, 22)
    - Body/bullets (paragraph 11-18)

    Then applies those exact formats to the new resume content.

    Args:
        content: Resume content in plain text
        output_path: Where to save the DOCX
        template_path: Path to Resume_AI.docx (optional)

    Returns:
        Path to created DOCX file
    """
    # Load Resume_AI.docx to extract formatting
    if template_path is None:
        template_path = Path(__file__).parent.parent.parent / "Resume_AI.docx"

    template_doc = Document(template_path)

    # Extract format templates from Resume_AI.docx
    format_templates = {
        'name': template_doc.paragraphs[0],           # Anmol Desai (20pt Bold)
        'contact': template_doc.paragraphs[1],        # Contact line (11pt)
        'section': template_doc.paragraphs[3],        # Summary (14pt Bold, 1.15)
        'job_title': template_doc.paragraphs[9],      # AI Engineer (10pt Bold)
        'company': template_doc.paragraphs[10],       # Omega Labs (10pt)
        'bullet': template_doc.paragraphs[11],        # Bullet point (10pt, 1.15)
        'body': template_doc.paragraphs[5],           # Body text (10pt, 1.15)
    }

    # Create new document
    doc = Document()

    # Parse content
    lines = content.split('\n')

    is_first_line = True
    in_contact = False

    for line in lines:
        line_stripped = line.strip()

        # Skip separators
        if (len(line_stripped) > 20 and
            (line_stripped.count('=') > 20 or
             line_stripped.count('-') > 20 or
             line_stripped.count('_') > 20)):
            continue

        # Empty lines
        if not line_stripped:
            doc.add_paragraph()
            continue

        # NAME - clone format from Resume_AI paragraph 0
        if is_first_line:
            p = doc.add_paragraph()
            clone_paragraph_format(format_templates['name'], p)
            run = p.add_run(line_stripped)
            if format_templates['name'].runs:
                clone_run_format(format_templates['name'].runs[0], run)
            is_first_line = False
            in_contact = True
            continue

        # CONTACT - clone format from Resume_AI paragraph 1
        if in_contact and ('@' in line_stripped or '|' in line_stripped or
                          'linkedin.com' in line_stripped.lower()):
            p = doc.add_paragraph()
            clone_paragraph_format(format_templates['contact'], p)
            run = p.add_run(line_stripped)
            if format_templates['contact'].runs:
                clone_run_format(format_templates['contact'].runs[0], run)
            in_contact = False
            continue

        # SECTION HEADERS - clone format from Resume_AI paragraph 3
        if (line_stripped.isupper() or
            any(kw in line_stripped.upper() for kw in
                ['EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS', 'SUMMARY',
                 'PROFESSIONAL', 'WORK', 'COMPETENCIES', 'TECHNICAL', 'CORE'])):
            p = doc.add_paragraph()
            clone_paragraph_format(format_templates['section'], p)
            run = p.add_run(line_stripped.replace('#', '').strip())
            if format_templates['section'].runs:
                clone_run_format(format_templates['section'].runs[0], run)
            continue

        # BULLET POINTS - clone format from Resume_AI paragraph 11
        if line_stripped.startswith('•') or line_stripped.startswith('-') or line_stripped.startswith('*'):
            bullet_text = line_stripped.lstrip('•-* ').strip()
            p = doc.add_paragraph()
            clone_paragraph_format(format_templates['bullet'], p)
            run = p.add_run(bullet_text)
            if format_templates['bullet'].runs:
                clone_run_format(format_templates['bullet'].runs[0], run)
            continue

        # JOB TITLES - clone format from Resume_AI paragraph 9
        if '|' in line_stripped or any(year in line_stripped for year in
                                       ['2020', '2021', '2022', '2023', '2024', '2025']):
            p = doc.add_paragraph()
            clone_paragraph_format(format_templates['job_title'], p)
            run = p.add_run(line_stripped)
            if format_templates['job_title'].runs:
                clone_run_format(format_templates['job_title'].runs[0], run)
            continue

        # REGULAR BODY TEXT - clone format from Resume_AI paragraph 5
        p = doc.add_paragraph()
        clone_paragraph_format(format_templates['body'], p)
        run = p.add_run(line_stripped)
        if format_templates['body'].runs:
            clone_run_format(format_templates['body'].runs[0], run)

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path


def create_cover_letter_docx(content: str, output_path: str, template_path: str = None) -> str:
    """Create cover letter DOCX."""
    doc = Document()

    lines = content.split('\n')

    for line in lines:
        line_stripped = line.strip()
        if line_stripped:
            p = doc.add_paragraph()
            run = p.add_run(line_stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path
